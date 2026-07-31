import os
from io import BytesIO

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from dotenv import load_dotenv

from career_core import (
    COMMON_SKIP_OPTIONS,
    CareerAnalyzer,
    ModelSettings,
    UserFacingError,
    build_base_resume,
    clean_text,
    extract_resume_file,
)


load_dotenv()

st.set_page_config(
    page_title="AI 求职助手",
    page_icon=":material/work:",
    layout="centered",
)


def get_setting(name: str, default: str = "") -> str:
    local = os.getenv(name)
    if local:
        return local
    try:
        value = st.secrets.get(name, default)
    except Exception:
        return default
    return str(value) if value else default


def get_analyzer() -> CareerAnalyzer:
    return CareerAnalyzer(
        ModelSettings(
            api_key=get_setting("DASHSCOPE_API_KEY"),
            workspace_id=get_setting("DASHSCOPE_WORKSPACE_ID"),
            model=get_setting("DASHSCOPE_MODEL", "qwen3.7-plus"),
            offline=get_setting("AI_RESUME_OFFLINE") == "1",
        )
    )


def initialize_state() -> None:
    defaults = {
        "view": "landing",
        "journey": None,
        "target_stage": "input",
        "target_diagnosis": None,
        "target_questions": [],
        "target_result": None,
        "direction_result": None,
        "base_result": None,
        "target_resume_input": "",
        "target_position": "",
        "target_jd": "",
        "target_answers": [],
        "notice": "",
        "error": "",
        "hide_city_advice": False,
    }
    for key, value in defaults.items():
        st.session_state.setdefault(key, value)


def clear_outputs() -> None:
    for key, value in {
        "target_stage": "input",
        "target_diagnosis": None,
        "target_questions": [],
        "target_result": None,
        "direction_result": None,
        "base_result": None,
        "target_answers": [],
        "notice": "",
        "error": "",
        "hide_city_advice": False,
    }.items():
        st.session_state[key] = value


def enter_journey(journey: str) -> None:
    clear_outputs()
    st.session_state.view = "workspace"
    st.session_state.journey = journey


def go_home() -> None:
    st.session_state.view = "landing"
    st.session_state.error = ""


def start_over() -> None:
    journey = st.session_state.journey
    clear_outputs()
    if journey != "target":
        st.session_state.target_resume_input = ""


def continue_base_to_target() -> None:
    base = st.session_state.get("base_result") or {}
    st.session_state.target_resume_input = base.get("plain_text", "")
    st.session_state.view = "workspace"
    st.session_state.journey = "target"
    st.session_state.target_stage = "input"
    st.session_state.error = ""


def continue_base_to_direction() -> None:
    base = st.session_state.get("base_result") or {}
    st.session_state.direction_experience = base.get("plain_text", "")
    st.session_state.view = "workspace"
    st.session_state.journey = "direction"
    st.session_state.direction_result = None
    st.session_state.error = ""


def build_word_document(markdown: str, title: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)

    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown.splitlines()
    if lines and lines[0].strip().lstrip("#").strip() == title:
        lines = lines[1:]

    for line in lines:
        text = line.strip()
        if not text:
            document.add_paragraph()
            continue
        if text.startswith("### "):
            document.add_heading(text[4:], level=2)
        elif text.startswith("## "):
            document.add_heading(text[3:], level=2)
        elif text.startswith("# "):
            document.add_heading(text[2:], level=1)
        elif text.startswith(("- ", "* ")):
            text = text[2:]
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(text.replace("**", "").replace("__", "").replace("`", ""))
        else:
            document.add_paragraph(text.replace("**", "").replace("__", "").replace("`", ""))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def markdown_to_plain_text(markdown: str) -> str:
    lines: list[str] = []
    for line in markdown.splitlines():
        text = line.strip()
        if text.startswith("# "):
            text = text[2:].strip()
        elif text.startswith("## "):
            text = text[3:].strip()
        elif text.startswith("### "):
            text = text[4:].strip()
        text = text.replace("**", "").replace("__", "").replace("`", "")
        lines.append(text)
    return "\n".join(lines).strip()


def build_resume_markdown(result: dict, position_name: str = "") -> str:
    bullets = result.get("resume_bullets", [])
    lines = [
        "# 岗位定制简历",
        "",
        "姓名：请补充    手机：请补充    邮箱：请补充    城市：请补充",
    ]
    if position_name:
        lines.extend(["", "## 求职目标", position_name])

    strengths = result.get("strengths", [])
    if strengths:
        lines.extend(["", "## 核心优势"])
        lines.extend(f"- {item}" for item in strengths[:4])

    lines.extend(["", "## 项目与经历亮点"])
    if bullets:
        lines.extend(f"- {item.get('text', '')}" for item in bullets)
    else:
        lines.append("- 当前事实不足，请先补充真实经历后再生成简历。")

    lines.extend(
        [
            "",
            "## 教育背景",
            "学校 / 专业 / 学历：请补充",
            "时间：请补充",
            "相关课程、证书或毕业设计：请补充",
            "",
            "## 技能与工具",
            "- 请补充真实掌握的工具、语言、平台或证书。",
        ]
    )

    unknowns = result.get("unknowns", [])
    if unknowns:
        lines.extend(["", "## 投递前待确认"])
        lines.extend(f"- {item}" for item in unknowns[:4])

    lines.extend(
        [
            "",
            "## 使用说明",
            "复制到简历后，请先补齐个人信息、教育背景和技能工具，再按真实情况删改。",
            "所有数据和职责应以真实经历为准，不能把课程项目写成企业项目。",
        ]
    )
    return "\n".join(lines).strip()


def build_base_resume_markdown(result: dict) -> str:
    lines = [
        "# 基础简历模板",
        "",
        "姓名：请补充    手机：请补充    邮箱：请补充    城市：请补充",
    ]
    title_map = {
        "教育背景": "教育背景",
        "实习、兼职或实践经历": "实践经历",
        "课程、个人或作品项目": "项目经历",
        "技能、工具与证书": "技能与证书",
        "可迁移技能": "可迁移能力",
    }
    for section in result.get("sections", []):
        title = title_map.get(section.get("title", ""), section.get("title", "经历"))
        content = clean_text(section.get("content"))
        if not content:
            continue
        lines.extend(["", f"## {title}", content])
    lines.extend(
        [
            "",
            "## 简历润色提醒",
            f"- 责任程度先按“{result.get('responsibility', '参与')}”记录，不夸大主导权。",
            "- 每段经历后续最好补充：任务、动作、工具、交付物、可验证结果。",
            "- 投递前再按具体岗位删除无关内容，保留最能证明匹配度的 3 到 5 条。",
        ]
    )
    return "\n".join(lines).strip()


def report_markdown(result: dict) -> str:
    lines = [
        "# 岗位匹配结论",
        result.get("match_conclusion", ""),
        f"硬性条件：{result.get('hard_requirements', '')}",
        f"结论可信度：{result.get('confidence', '低')}",
        "",
        "# 招聘方真正看重什么",
    ]
    lines.extend(f"- {item}" for item in result.get("requirements", []))
    lines.extend(["", "# 最可能影响初筛的风险"])
    lines.extend(f"- {item}" for item in result.get("risks", []))
    lines.extend(["", "# 已有的真实优势"])
    lines.extend(f"- {item}" for item in result.get("strengths", []))
    lines.extend(["", "# 修改后的定制简历"])
    lines.extend(f"- {item.get('text', '')}" for item in result.get("resume_bullets", []))
    lines.extend(["", "# 仍需确认的信息"])
    lines.extend(f"- {item}" for item in result.get("unknowns", []))
    lines.extend(["", "# 面试官可能追问"])
    lines.extend(f"- {item}" for item in result.get("interview_questions", []))
    return "\n".join(lines).strip()


def render_landing() -> None:
    st.title("AI 求职助手")
    st.caption("先看清岗位，再整理证据，最后生成可解释、可复核的求职材料。")
    st.subheader("你现在最需要完成哪件事？")
    st.caption("首页只有三个入口。选择最接近你当前情况的一项即可。")

    with st.container(border=True):
        st.badge("有 JD 直接选这里", icon=":material/bolt:", color="blue")
        st.markdown("#### :material/target: 我有目标岗位，准备投递")
        st.caption(
            "已有岗位名称或招聘要求，直接分析匹配情况并生成定制简历，不需要填写职业偏好问卷。"
        )
        st.markdown("适合：已有真实 JD、已知道岗位名称、当前要修改简历并准备投递。")
        st.button(
            "开始岗位定制",
            icon=":material/arrow_forward:",
            type="primary",
            width="stretch",
            on_click=enter_journey,
            args=("target",),
            key="landing_target",
        )

    with st.container(border=True):
        st.markdown("#### :material/explore: 我还没确定投什么岗位")
        st.caption(
            "结合你的经历、工作偏好和城市考虑，梳理可以直接尝试、相邻转向和继续探索的职业方向。"
        )
        st.markdown("适合：方向不清、想转行、有几个方向难选择，或需要城市参考。")
        st.button(
            "探索岗位方向",
            icon=":material/arrow_forward:",
            width="stretch",
            on_click=enter_journey,
            args=("direction",),
            key="landing_direction",
        )

    with st.container(border=True):
        st.markdown("#### :material/article: 我想先整理一份基础简历")
        st.caption(
            "没有完整简历也没关系，从课程、项目、实习、兼职和作品中整理出一份真实可用的基础版本。"
        )
        st.markdown("适合：没有完整简历、经历比较零散，或想先把真实经历整理清楚。")
        st.button(
            "整理基础简历",
            icon=":material/arrow_forward:",
            width="stretch",
            on_click=enter_journey,
            args=("base",),
            key="landing_base",
        )


def render_diagnosis(diagnosis: dict) -> None:
    st.subheader("初步诊断")
    st.caption("先看已有信息能说明什么，再决定是否补充。此时还没有生成定制简历。")
    if diagnosis.get("notice"):
        st.warning(diagnosis["notice"], icon=":material/info:")
    with st.container(border=True):
        st.markdown("#### 招聘方真正看重的要求")
        for item in diagnosis.get("requirements", []):
            st.markdown(f"- {item}")
    with st.container(border=True):
        st.markdown("#### 你已经能证明的能力")
        for item in diagnosis.get("evidence", []):
            st.markdown(f"- {item}")
    with st.container(border=True):
        st.markdown("#### 最可能影响初筛的 3 个风险")
        for item in diagnosis.get("risks", [])[:3]:
            st.markdown(f"- {item}")
    with st.container(border=True):
        st.markdown("#### 暂时无法判断")
        for item in diagnosis.get("unknowns", []):
            st.markdown(f"- {item}")
    st.info(diagnosis.get("priority", ""), icon=":material/edit_note:")


def collect_target_answers(questions: list[dict]) -> list[dict[str, str]]:
    answers: list[dict[str, str]] = []
    for index, question in enumerate(questions):
        selected = st.session_state.get(f"target_q_{index}", "先按现有内容生成")
        detail = clean_text(st.session_state.get(f"target_q_detail_{index}", ""))
        answer = detail or selected
        answers.append(
            {
                "id": question.get("id", f"q{index + 1}"),
                "question": question.get("question", ""),
                "answer": answer,
                "status": (
                    "skipped" if selected in COMMON_SKIP_OPTIONS and not detail else "confirmed"
                ),
            }
        )
    return answers


def finalize_target(answers: list[dict[str, str]]) -> None:
    context = st.session_state.target_context
    analyzer = get_analyzer()
    with st.spinner("正在进行 STAR 定向重构和独立真实性复核……"):
        result = analyzer.finalize(
            context["position"],
            context["jd"],
            context["resume"],
            st.session_state.target_diagnosis,
            answers,
        )
    st.session_state.target_answers = answers
    st.session_state.target_result = result
    st.session_state.target_stage = "result"
    if result.get("offline"):
        st.session_state.notice = (
            "当前版本仅使用已确认信息，没有添加未经确认的数据。"
            "补充更多事实或恢复 AI 服务后，可以进一步增强。"
        )
    st.rerun()


def render_target_flow() -> None:
    st.title("岗位定制")
    st.caption("已有明确目标岗位的用户只走这条流程，不需要填写职业偏好或城市问卷。")

    if st.session_state.target_stage == "input":
        st.info(
            "建议上传前隐藏身份证号、完整住址等非必要敏感信息。"
            "测试版本不应长期保存简历正文。",
            icon=":material/privacy_tip:",
        )
        with st.form("target_input_form"):
            position = st.text_input(
                "目标岗位名称",
                key="target_position",
                placeholder="例如：新媒体运营专员",
                persist_state="session",
            )
            jd = st.text_area(
                "真实岗位 JD",
                key="target_jd",
                placeholder="粘贴岗位职责、任职要求和硬性条件。内容过短会影响判断。",
                height=220,
                persist_state="session",
            )
            upload = st.file_uploader(
                "上传简历（可选）",
                type=["pdf", "docx", "txt", "md"],
                key="target_resume_upload",
                help="支持 PDF、Word（.docx）、TXT 和 Markdown；也可以直接粘贴经历。",
            )
            resume_input = st.text_area(
                "或粘贴真实经历",
                key="target_resume_input",
                placeholder="教育、实习、兼职、课程项目、个人项目、社团、作品和会用的工具都可以。",
                height=260,
                persist_state="session",
            )
            submitted = st.form_submit_button(
                "查看初步诊断",
                icon=":material/analytics:",
                type="primary",
                width="stretch",
            )

        if submitted:
            st.session_state.error = ""
            try:
                uploaded_text = ""
                if upload is not None:
                    uploaded_text = extract_resume_file(upload.name, upload.getvalue())
                resume = "\n\n".join(
                    item for item in (uploaded_text, clean_text(resume_input)) if item
                )
                if not clean_text(position):
                    raise UserFacingError("请填写目标岗位名称。")
                if len(clean_text(jd)) < 50:
                    raise UserFacingError(
                        "岗位 JD 过短，至少粘贴一段岗位职责或任职要求后再分析。"
                    )
                if not resume:
                    raise UserFacingError("简历正文为空，请上传可读取文件或粘贴真实经历。")
                analyzer = get_analyzer()
                with st.spinner("正在拆解 JD、提取真实证据并识别缺口……"):
                    diagnosis = analyzer.diagnose(position, jd, resume)
                    questions = analyzer.questions(diagnosis, resume)
                st.session_state.target_context = {
                    "position": clean_text(position),
                    "jd": clean_text(jd),
                    "resume": resume,
                }
                st.session_state.target_diagnosis = diagnosis
                st.session_state.target_questions = questions[:3]
                st.session_state.target_stage = "diagnosis"
                st.rerun()
            except UserFacingError as exc:
                st.session_state.error = str(exc)

    elif st.session_state.target_stage == "diagnosis":
        render_diagnosis(st.session_state.target_diagnosis)
        questions = st.session_state.target_questions[:3]
        st.subheader("补充最多 3 个关键事实")
        st.caption("每题都有常见选项，也可以不记得、没有数据、跳过或直接按现有内容生成。")
        with st.form("target_questions_form"):
            for index, question in enumerate(questions):
                with st.container(border=True):
                    st.markdown(f"**{index + 1}. {question['question']}**")
                    options = question.get("choices", []) + list(COMMON_SKIP_OPTIONS)
                    st.radio(
                        f"第 {index + 1} 题选项",
                        options,
                        index=len(question.get("choices", [])) + 3,
                        key=f"target_q_{index}",
                        label_visibility="collapsed",
                    )
                    st.text_input(
                        "补充一句真实细节（选填）",
                        key=f"target_q_detail_{index}",
                        placeholder="例如：参与 5 人团队；约一个月完成；根据反馈修改 2 至 3 轮。",
                    )
            answer_submit = st.form_submit_button(
                "根据回答生成定制结果",
                icon=":material/arrow_forward:",
                type="primary",
                width="stretch",
            )
        if answer_submit:
            finalize_target(collect_target_answers(questions))
        if st.button(
            "跳过全部，按现有内容生成保守版本",
            icon=":material/fast_forward:",
            width="stretch",
            key="skip_all_target_questions",
        ):
            answers = [
                {
                    "id": question.get("id", ""),
                    "question": question.get("question", ""),
                    "answer": "先按现有内容生成",
                    "status": "skipped",
                }
                for question in questions
            ]
            finalize_target(answers)
        if st.button(
            "返回修改岗位或经历",
            icon=":material/arrow_back:",
            width="stretch",
            key="back_to_target_input",
        ):
            st.session_state.target_stage = "input"
            st.rerun()
    else:
        render_target_result()

    if st.session_state.error:
        st.error(st.session_state.error, icon=":material/error:")


def render_target_result() -> None:
    result = st.session_state.target_result or {}
    context = st.session_state.get("target_context") or {}
    position_name = context.get("position", "")
    st.title("岗位定向结果")
    if st.session_state.notice:
        st.warning(st.session_state.notice, icon=":material/verified_user:")

    st.subheader("1. 岗位匹配结论")
    with st.container(border=True):
        st.markdown(result.get("match_conclusion", "当前信息不足，暂时无法形成可靠结论。"))
        st.markdown(f"**硬性条件：** {result.get('hard_requirements', '仍需核对')}")
        st.markdown(f"**结论可信度：** {result.get('confidence', '低')}")

    st.subheader("2. 招聘方真正看重什么")
    for item in result.get("requirements", []):
        with st.container(border=True):
            st.markdown(item)

    st.subheader("3. 最可能影响初筛的 3 个风险")
    for item in result.get("risks", [])[:3]:
        st.warning(item, icon=":material/warning:")

    st.subheader("4. 你已有的真实优势")
    for item in result.get("strengths", []):
        st.success(item, icon=":material/check_circle:")

    st.subheader("5. 修改前后对照")
    before_after = result.get("before_after", [])
    if before_after:
        for item in before_after:
            with st.container(border=True):
                st.caption(f"原文：{item.get('before', '未提供')}")
                st.markdown(f"**改写：{item.get('after', '')}**")
                st.caption(f"依据：{item.get('basis', '已确认事实')}")
    else:
        st.caption("保守版本没有强行改写原文；补充真实责任和交付后可生成更清晰的对照。")

    st.subheader("6. 修改后的定制简历")
    bullets = result.get("resume_bullets", [])
    for index, item in enumerate(bullets):
        with st.container(border=True):
            st.markdown(f"**{item.get('text', '')}**")
            st.caption(f"事实状态：{item.get('fact_status', '已确认事实')}")
            with st.popover("纠正这一条", icon=":material/edit:"):
                correction = st.selectbox(
                    "选择纠正方式",
                    [
                        "这不是我负责的",
                        "数据不准确",
                        "表达太夸张",
                        "表达太笼统",
                        "不像我的说话方式",
                        "保留原文",
                        "重新改写这一条",
                    ],
                    key=f"correction_type_{index}",
                )
                detail = st.text_input(
                    "补充说明（选填）",
                    key=f"correction_detail_{index}",
                    placeholder="只写需要更正的真实情况。",
                )
                if st.button(
                    "只更新这一条",
                    key=f"apply_correction_{index}",
                    type="primary",
                    width="stretch",
                ):
                    facts = st.session_state.target_context["resume"] + "\n" + "\n".join(
                        answer.get("answer", "")
                        for answer in st.session_state.target_answers
                    )
                    feedback = "；".join(part for part in (correction, detail) if part)
                    revised = get_analyzer().revise(facts, item.get("text", ""), feedback)
                    st.session_state.target_result["resume_bullets"][index]["text"] = revised[
                        "replacement"
                    ]
                    st.session_state.target_result["resume_bullets"][index][
                        "fact_status"
                    ] = revised["status"]
                    st.session_state.notice = revised["basis"]
                    st.rerun()

    review = result.get("review", {})
    with st.container(border=True):
        st.markdown("#### 独立真实性复核")
        st.badge(
            review.get("status", "已复核"),
            icon=":material/verified:",
            color="green",
        )
        for note in review.get("notes", []):
            st.markdown(f"- {note}")
        rejected = review.get("rejected_claims", [])
        if rejected:
            st.caption(f"已拦截 {len(rejected)} 条缺少事实支撑的表述。")

    st.subheader("7. 仍需确认的信息")
    for item in result.get("unknowns", []):
        st.info(item, icon=":material/help:")

    st.subheader("8. 面试官可能追问")
    for index, item in enumerate(result.get("interview_questions", []), start=1):
        with st.container(border=True):
            st.markdown(f"**{index}. {item}**")

    markdown = report_markdown(result)
    resume_markdown = build_resume_markdown(result, position_name)
    resume_text = markdown_to_plain_text(resume_markdown)
    st.subheader("9. 下载材料")
    with st.container(border=True):
        st.markdown("#### 可复制简历正文")
        st.caption("这里是可以直接粘到简历里的文本；方括号和“请补充”内容投递前需要替换。")
        st.text_area(
            "复制这份定制简历文本",
            value=resume_text,
            height=420,
            key="target_resume_copy_text",
        )
        with st.container(horizontal=True, horizontal_alignment="distribute"):
            st.download_button(
                "下载有内容的简历模板（Word）",
                build_word_document(resume_markdown, "岗位定制简历"),
                file_name="岗位定制简历.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                icon=":material/download:",
                width="stretch",
                key="download_resume",
            )
            st.download_button(
                "下载分析报告（Word）",
                build_word_document(markdown, "AI 岗位定向分析报告"),
                file_name="AI岗位定向分析报告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                icon=":material/download:",
                width="stretch",
            )


def render_base_flow() -> None:
    st.title("整理基础简历")
    st.caption("没有正式工作经历也可以填写课程、个人项目、比赛、社团、兼职、作品和长期负责的事务。")
    if not st.session_state.base_result:
        with st.form("base_resume_form"):
            education = st.text_area(
                "教育背景",
                key="base_education",
                placeholder="学校、专业、时间、相关课程或毕业设计。",
                height=110,
                persist_state="session",
            )
            experience = st.text_area(
                "实习、兼职、社团或实践经历",
                key="base_experience",
                placeholder="做过什么、你负责哪部分、最终完成了什么；没有可以留空。",
                height=140,
                persist_state="session",
            )
            projects = st.text_area(
                "课程、个人、练习或作品项目",
                key="base_projects",
                placeholder="请保留真实标签，不要把课程或个人项目写成企业项目。",
                height=140,
                persist_state="session",
            )
            skills = st.text_area(
                "技能、工具与证书",
                key="base_skills",
                placeholder="只写真实会用的工具、语言和证书。",
                height=110,
                persist_state="session",
            )
            transferable = st.text_area(
                "可迁移技能（选填）",
                key="base_transferable",
                placeholder="例如：资料整理、沟通、测试、内容制作、活动协作。",
                height=100,
                persist_state="session",
            )
            responsibility = st.segmented_control(
                "这段主要经历中的责任程度",
                ["独立完成", "主导", "参与", "协助"],
                default="参与",
                key="base_responsibility",
            )
            submitted = st.form_submit_button(
                "生成真实基础简历",
                icon=":material/article:",
                type="primary",
                width="stretch",
            )
        if submitted:
            try:
                st.session_state.base_result = build_base_resume(
                    {
                        "education": education,
                        "experience": experience,
                        "projects": projects,
                        "skills": skills,
                        "transferable": transferable,
                        "responsibility": responsibility,
                    }
                )
                st.rerun()
            except UserFacingError as exc:
                st.session_state.error = str(exc)
    else:
        result = st.session_state.base_result
        base_markdown = build_base_resume_markdown(result)
        base_text = markdown_to_plain_text(base_markdown)
        st.success(result["notice"], icon=":material/verified:")
        st.text_area(
            "可复制的基础简历模板",
            value=base_text,
            height=480,
            key="base_resume_preview",
        )
        st.download_button(
            "下载有内容的基础简历模板（Word）",
            build_word_document(base_markdown, "基础简历"),
            file_name="基础简历.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            icon=":material/download:",
            width="stretch",
        )
        st.subheader("下一步")
        st.button(
            "我有目标岗位，继续定制",
            icon=":material/target:",
            type="primary",
            width="stretch",
            on_click=continue_base_to_target,
            key="base_to_target",
        )
        st.button(
            "我还没确定方向，帮我推荐",
            icon=":material/explore:",
            width="stretch",
            on_click=continue_base_to_direction,
            key="base_to_direction",
        )
    if st.session_state.error:
        st.error(st.session_state.error, icon=":material/error:")


def render_direction_flow() -> None:
    st.title("探索岗位方向")
    st.caption("第一轮只判断 3 件必要的事。过去经历只作为证据，不替你决定未来。")
    if not st.session_state.direction_result:
        with st.form("direction_form"):
            experience = st.text_area(
                "先写下真实经历或基础简历",
                key="direction_experience",
                placeholder="课程、项目、实习、兼职、社团、作品、工具和长期负责的事情都可以。",
                height=220,
                persist_state="session",
            )
            continue_old = st.segmented_control(
                "1. 过去做过的方向还想继续吗？",
                ["想继续", "不想继续", "不确定", "跳过"],
                default="不确定",
                key="direction_continue_old",
            )
            priority = st.segmented_control(
                "2. 岗位和城市，哪个更优先？",
                ["岗位优先", "城市优先", "暂时没有偏好", "跳过"],
                default="暂时没有偏好",
                key="direction_priority",
            )
            excluded_choice = st.multiselect(
                "3. 明确不接受哪些工作内容？",
                ["纯销售", "高频出差", "夜班", "大量电话沟通", "长期加班", "暂时没有偏好", "跳过"],
                default=["暂时没有偏好"],
                key="direction_excluded_choice",
            )
            excluded_detail = st.text_input(
                "其他不接受内容（选填）",
                key="direction_excluded_detail",
                placeholder="例如：不想继续游戏开发；不考虑纯客服。",
            )
            urgency = st.segmented_control(
                "是否急需尽快就业？（选填）",
                ["是", "否", "不确定", "跳过"],
                default="不确定",
                key="direction_urgency",
            )
            submitted = st.form_submit_button(
                "生成方向建议",
                icon=":material/explore:",
                type="primary",
                width="stretch",
            )
        if submitted:
            excluded = "；".join(
                item
                for item in excluded_choice + [clean_text(excluded_detail)]
                if item and item not in {"暂时没有偏好", "跳过"}
            )
            profile = {
                "experience": clean_text(experience),
                "continue_old": continue_old or "不确定",
                "priority": priority or "暂时没有偏好",
                "excluded": excluded or "暂时没有偏好",
                "urgency": urgency or "不确定",
            }
            with st.spinner("正在拆解可迁移能力并梳理三类方向……"):
                st.session_state.direction_result = get_analyzer().direction(profile)
            st.session_state.direction_profile = profile
            st.rerun()
    else:
        result = st.session_state.direction_result
        if result.get("offline"):
            st.warning(result.get("notice", ""), icon=":material/info:")
        else:
            st.caption(result.get("notice", ""))
        for role_index, role in enumerate(result.get("roles", [])[:3]):
            with st.container(border=True):
                st.badge(role.get("category", "探索方向"), color="blue")
                st.subheader(role.get("title", "待验证方向"))
                st.markdown(f"**日常实际做什么：** {role.get('daily_work', '')}")
                st.markdown(f"**为什么推荐：** {role.get('why', '')}")
                st.markdown(f"**真实依据：** {role.get('evidence', '')}")
                st.markdown(f"**已经具备：** {role.get('existing', '')}")
                st.markdown(f"**主要缺口：** {role.get('gap', '')}")
                st.markdown(f"**可能不符合偏好：** {role.get('preference_risk', '')}")
                st.markdown(f"**现在是否建议投递：** {role.get('apply_now', '')}")
                st.markdown(f"**下一步：** {role.get('next_action', '')}")
                if st.button(
                    "我不想继续这个方向",
                    icon=":material/close:",
                    key=f"remove_direction_{role_index}",
                    width="stretch",
                ):
                    st.session_state.direction_result["roles"].pop(role_index)
                    st.rerun()

        if not st.session_state.hide_city_advice:
            city = result.get("cities", {})
            with st.container(border=True):
                st.subheader("城市第一梯队")
                st.markdown(f"**{city.get('tier', '')}**")
                st.markdown(city.get("summary", ""))
                roles = "、".join(city.get("search_roles", []))
                if roles:
                    st.markdown(f"**可以搜索：** {roles}")
                st.caption(city.get("notice", ""))
                if st.button(
                    "不考虑这些城市",
                    icon=":material/location_off:",
                    width="stretch",
                    key="hide_city",
                ):
                    st.session_state.hide_city_advice = True
                    st.rerun()

        micro = result.get("micro_project", {})
        with st.container(border=True):
            st.subheader("经历证据较少时：先做一个真实微型项目")
            st.markdown(f"**周期：** {micro.get('duration', '3 至 7 天')}")
            st.markdown(f"**任务：** {micro.get('task', '')}")
            st.markdown(f"**最终交付：** {micro.get('deliverable', '')}")
            st.markdown(f"**完成后如何写：** {micro.get('resume_usage', '')}")
    if st.session_state.error:
        st.error(st.session_state.error, icon=":material/error:")


initialize_state()

if st.session_state.view == "landing":
    render_landing()
    st.stop()

st.button(
    "返回首页",
    icon=":material/arrow_back:",
    width="stretch",
    on_click=go_home,
    key="global_back_home",
)
if (
    st.session_state.target_result
    or st.session_state.direction_result
    or st.session_state.base_result
):
    st.button(
        "开始新的分析",
        icon=":material/restart_alt:",
        width="stretch",
        on_click=start_over,
        key="global_start_over",
    )

journey = st.session_state.journey
if journey == "target":
    render_target_flow()
elif journey == "base":
    render_base_flow()
else:
    render_direction_flow()
